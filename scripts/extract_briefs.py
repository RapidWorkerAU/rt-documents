"""
extract_briefs.py

Reads the HSESC Management System Playbook (playbook.docx) and extracts
one structured JSON brief per procedure. Each brief is saved to the briefs/
directory as PRO-XXXX.json.

The brief JSON drives the document generation step — it contains:
  - procedure_id, title, document_number, management_system_element
  - owns: list of content items this document owns
  - references: list of documents this procedure references
  - referenced_by: list of documents that reference this procedure
  - preamble: the purpose paragraph for this procedure
  - requirements: list of {tag, source, text} requirement objects
  - project_completion_guidance: list of PCG numbered items
  - external_sources: list of {framework, clause, summary} traceability items
"""

import argparse
import base64
import json
import os
import re
import sys
import time

import anthropic
from docx import Document


# ── Procedure registry ────────────────────────────────────────────────────────
# All 51 procedures in document order. This is used to chain jobs and to
# validate that every brief has been extracted.

PROCEDURE_REGISTRY = [
    {"id": "PRO-0001", "title": "Legal and Other Requirements Procedure"},
    {"id": "PRO-0002", "title": "HSESC Risk Management Procedure"},
    {"id": "PRO-0003", "title": "Organisational Resources and Responsibilities Procedure"},
    {"id": "PRO-0004", "title": "Training, Competency and Awareness Procedure"},
    {"id": "PRO-0005", "title": "Supplier and Contractor Management Procedure"},
    {"id": "PRO-0006", "title": "Fitness for Work Procedure"},
    {"id": "PRO-0007", "title": "Management Improvement Planning Procedure"},
    {"id": "PRO-0008", "title": "Working at Heights Procedure"},
    {"id": "PRO-0009", "title": "Incident and Action Management Procedure"},
    {"id": "PRO-0010", "title": "Documentation and Document Control Procedure"},
    {"id": "PRO-0011", "title": "Communication and Consultation Procedure"},
    {"id": "PRO-0012", "title": "Chemical and Hazardous Substances Procedure"},
    {"id": "PRO-0013", "title": "Noise Exposure Control Procedure"},
    {"id": "PRO-0014", "title": "Confined Spaces Procedure"},
    {"id": "PRO-0015", "title": "Cranes and Lifting Procedure"},
    {"id": "PRO-0016", "title": "Hot Works Procedure"},
    {"id": "PRO-0017", "title": "Manual Tasks and Ergonomics Procedure"},
    {"id": "PRO-0018", "title": "Aviation Procedure"},
    {"id": "PRO-0019", "title": "Process Safety Procedure"},
    {"id": "PRO-0020", "title": "Functional Safety Procedure"},
    {"id": "PRO-0021", "title": "Permit to Work Procedure"},
    {"id": "PRO-0022", "title": "Infectious Disease Control Procedure"},
    {"id": "PRO-0023", "title": "Radiation Control Procedure"},
    {"id": "PRO-0024", "title": "Isolation Procedure"},
    {"id": "PRO-0025", "title": "Electrical Safety Procedure"},
    {"id": "PRO-0026", "title": "Vehicles and Driving Procedure"},
    {"id": "PRO-0027", "title": "Personal Protective Equipment Procedure"},
    {"id": "PRO-0028", "title": "Barricading, Signage and Demarcation Procedure"},
    {"id": "PRO-0029", "title": "Excavation and Penetration Procedure"},
    {"id": "PRO-0030", "title": "Working Over Water and Marine Vetting Activities Procedure"},
    {"id": "PRO-0031", "title": "Traffic Management Procedure"},
    {"id": "PRO-0032", "title": "Tools and Equipment Procedure"},
    {"id": "PRO-0033", "title": "Security Management Procedure"},
    {"id": "PRO-0034", "title": "Water Quality Protection and Water Management Procedure"},
    {"id": "PRO-0035", "title": "Air Quality Procedure"},
    {"id": "PRO-0036", "title": "Chemically Reactive Mineral Waste Control Procedure"},
    {"id": "PRO-0037", "title": "Land Disturbance Control and Rehabilitation Procedure"},
    {"id": "PRO-0038", "title": "Biodiversity Procedure"},
    {"id": "PRO-0039", "title": "Hazardous Materials and Non-Mineral Waste Control Procedure"},
    {"id": "PRO-0040", "title": "Management of Change Procedure"},
    {"id": "PRO-0041", "title": "Business Resilience and Recovery Procedure"},
    {"id": "PRO-0042", "title": "Monitoring and Measuring Procedure"},
    {"id": "PRO-0043", "title": "Data and Records Management Procedure"},
    {"id": "PRO-0044", "title": "Performance Assessment and Auditing Procedure"},
    {"id": "PRO-0045", "title": "Management Review Procedure"},
    {"id": "PRO-0046", "title": "Scaffolding Procedure"},
    {"id": "PRO-0047", "title": "Commissioning Safety Procedure"},
    {"id": "PRO-0048", "title": "Simultaneous Operations (SIMOPS) Procedure"},
    {"id": "PRO-0049", "title": "Communities and Social Performance Procedure"},
    {"id": "PRO-0050", "title": "Leadership Development Programme Procedure"},
    {"id": "PRO-0051", "title": "Explosives Management Procedure"},
]


def read_docx_text(path: str) -> str:
    """Extract all text from a docx file as a single string."""
    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Also extract table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def docx_to_base64(path: str) -> str:
    """Read a docx file and return base64-encoded bytes."""
    with open(path, "rb") as f:
        return base64.standard_b64encode(f.read()).decode("utf-8")


def extract_brief_for_procedure(
    client: anthropic.Anthropic,
    playbook_text: str,
    procedure: dict,
) -> dict:
    """
    Call Claude to extract the structured brief for one procedure from the playbook text.
    Returns a dict matching the brief JSON schema.
    """
    proc_id = procedure["id"]
    proc_title = procedure["title"]

    system_prompt = """You are a document engineering assistant working on a Rio Tinto owner-led project 
HSESC Management System. You are reading the HSESC Management System Playbook and extracting 
structured procedure briefs from it.

Your output must be valid JSON only. No markdown, no code fences, no preamble. 
Return only the JSON object described in the user message."""

    user_prompt = f"""Extract the complete procedure brief for {proc_id} ({proc_title}) from the playbook text below.

Return a JSON object with exactly this structure:

{{
  "procedure_id": "{proc_id}",
  "document_number": "RTPR-HSE-PRO-{proc_id[4:]}",
  "title": "{proc_title}",
  "management_system_element": "<element number and name from the brief, e.g. '02 - Legal & Other Requirements'>",
  "preamble": "<the purpose paragraph for this procedure in the owner-led project context>",
  "owns": [
    "<each item this document owns — verbatim from the Document Relationships block>"
  ],
  "references": [
    "<each document this procedure references — verbatim from the Document Relationships block>"
  ],
  "referenced_by": [
    "<each document that references this procedure — verbatim from the Document Relationships block>"
  ],
  "requirements": [
    {{
      "section_heading": "<the heading this requirement sits under in the brief>",
      "tag": "<the APPC or RT source tag, e.g. '[APPC – Appendix C, C-1 §19.0.1]' or '[RT Standard]'>",
      "text": "<the full requirement text verbatim>"
    }}
  ],
  "project_completion_guidance": [
    "<each numbered PCG item verbatim — these are the amber-coloured items in the brief>"
  ],
  "external_sources": [
    {{
      "framework": "<e.g. Appendix C, RT Standard, IFC PS1, ICMM>",
      "clause": "<specific clause reference, e.g. C-1 §19.0.1>",
      "summary": "<one sentence describing what this source requires>"
    }}
  ]
}}

Rules:
- Extract ONLY the content for {proc_id}. Do not mix in content from other procedures.
- Preserve all requirement text verbatim — do not paraphrase or summarise.
- For requirements without an APPC tag, use "[RT Standard]" as the tag.
- The external_sources list must include one entry per unique source clause referenced in the requirements.
- If a field has no content in the brief, use an empty array [] or empty string "".
- Return only valid JSON. No markdown, no code fences, no commentary.

PLAYBOOK TEXT:
{playbook_text}"""

    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=8000,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}],
            )

            raw = response.content[0].text.strip()

            # Strip any accidental markdown fences
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)

            brief = json.loads(raw)

            # Validate required fields
            required = [
                "procedure_id", "document_number", "title",
                "management_system_element", "preamble",
                "owns", "references", "referenced_by",
                "requirements", "project_completion_guidance", "external_sources"
            ]
            for field in required:
                if field not in brief:
                    raise ValueError(f"Missing required field: {field}")

            return brief

        except json.JSONDecodeError as e:
            print(f"  JSON parse error on attempt {attempt + 1}: {e}")
            if attempt < max_retries - 1:
                time.sleep(5)
            else:
                raise
        except Exception as e:
            print(f"  Error on attempt {attempt + 1}: {e}")
            if attempt < max_retries - 1:
                time.sleep(10)
            else:
                raise


def main():
    parser = argparse.ArgumentParser(description="Extract procedure briefs from playbook")
    parser.add_argument("--playbook", required=True, help="Path to playbook.docx")
    parser.add_argument("--output-dir", required=True, help="Directory to write brief JSON files")
    parser.add_argument(
        "--procedure",
        help="Extract only this procedure ID (e.g. PRO-0001). Default: all.",
        default=None,
    )
    args = parser.parse_args()

    # Validate inputs
    if not os.path.exists(args.playbook):
        print(f"ERROR: Playbook not found: {args.playbook}")
        sys.exit(1)

    os.makedirs(args.output_dir, exist_ok=True)

    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    print(f"Reading playbook: {args.playbook}")
    playbook_text = read_docx_text(args.playbook)
    print(f"Playbook text length: {len(playbook_text):,} characters")

    # Determine which procedures to extract
    if args.procedure:
        procedures = [p for p in PROCEDURE_REGISTRY if p["id"] == args.procedure]
        if not procedures:
            print(f"ERROR: Unknown procedure ID: {args.procedure}")
            sys.exit(1)
    else:
        procedures = PROCEDURE_REGISTRY

    print(f"Extracting {len(procedures)} procedure brief(s)...")

    success_count = 0
    fail_count = 0

    for procedure in procedures:
        proc_id = procedure["id"]
        output_path = os.path.join(args.output_dir, f"{proc_id}.json")

        # Skip if already extracted (allows resume)
        if os.path.exists(output_path):
            print(f"  SKIP {proc_id} — brief already exists")
            success_count += 1
            continue

        print(f"  Extracting {proc_id}: {procedure['title']}...")

        try:
            brief = extract_brief_for_procedure(client, playbook_text, procedure)

            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(brief, f, indent=2, ensure_ascii=False)

            req_count = len(brief.get("requirements", []))
            pcg_count = len(brief.get("project_completion_guidance", []))
            print(f"    OK — {req_count} requirements, {pcg_count} PCG items")
            success_count += 1

        except Exception as e:
            print(f"    FAILED: {e}")
            fail_count += 1
            # Write a minimal placeholder so the pipeline can continue
            placeholder = {
                "procedure_id": proc_id,
                "document_number": f"RTPR-HSE-PRO-{proc_id[4:]}",
                "title": procedure["title"],
                "management_system_element": "",
                "preamble": "",
                "owns": [],
                "references": [],
                "referenced_by": [],
                "requirements": [],
                "project_completion_guidance": [],
                "external_sources": [],
                "_extraction_failed": True,
                "_error": str(e),
            }
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(placeholder, f, indent=2, ensure_ascii=False)

        # Rate limit courtesy pause
        time.sleep(2)

    print(f"\nExtraction complete: {success_count} succeeded, {fail_count} failed")

    # Write the registry JSON for use by other scripts
    registry_path = os.path.join(args.output_dir, "_registry.json")
    with open(registry_path, "w", encoding="utf-8") as f:
        json.dump(PROCEDURE_REGISTRY, f, indent=2)
    print(f"Registry written: {registry_path}")

    if fail_count > 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
