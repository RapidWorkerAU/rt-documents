"""
generate_procedure.py

Generates a procedure document by:
1. Copying the procedure template (inputs/procedure_template.docx) — preserving all
   styles, headers, footers, logos, page numbers, and RT branding exactly.
2. Replacing PROCEDURE TITLE with the actual procedure title.
3. Replacing the body content placeholder with Claude-generated content.

This means every generated document is on the actual RT template automatically.
"""

import argparse
import json
import os
import re
import shutil
import sys
import time
import zipfile
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


# ── Colours ───────────────────────────────────────────────────────────────────
COLOUR_BLACK   = RGBColor(0x2D, 0x2D, 0x2D)
COLOUR_BLUE    = RGBColor(0x00, 0x59, 0xD1)
COLOUR_ORANGE  = RGBColor(0xC5, 0x5A, 0x11)
COLOUR_GREEN   = RGBColor(0x37, 0x56, 0x23)
COLOUR_HEADING = RGBColor(0x1F, 0x38, 0x64)
COLOUR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

# Word XML namespace
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── XML paragraph builders ────────────────────────────────────────────────────
# These build raw lxml elements that can be inserted directly into document.xml

def xml_para(style=None, space_before=None, space_after=None):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    if style:
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
    if space_before is not None or space_after is not None:
        spacing = OxmlElement("w:spacing")
        if space_before is not None:
            spacing.set(qn("w:before"), str(space_before))
        if space_after is not None:
            spacing.set(qn("w:after"), str(space_after))
        pPr.append(spacing)
    p.append(pPr)
    return p


def xml_run(para, text, bold=False, colour_hex=None, size_pt=10, italic=False):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), "Arial")
    rPr.append(fonts)

    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if colour_hex:
        col = OxmlElement("w:color")
        col.set(qn("w:val"), colour_hex.lstrip("#"))
        rPr.append(col)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    rPr.append(szCs)

    r.append(rPr)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    para.append(r)
    return r


def xml_border_bottom(para, colour="2E75B6", size=6):
    pPr = para.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para.insert(0, pPr)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), colour)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg_xml(cell_elem, hex_colour):
    tcPr = cell_elem.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell_elem.insert(0, tcPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


# ── High-level element builders ───────────────────────────────────────────────

def build_h1(number, title):
    p = xml_para(space_before=240, space_after=60)
    xml_border_bottom(p)
    xml_run(p, f"{number}  {title}", bold=True, colour_hex="1F3864", size_pt=13)
    return [p, xml_blank()]


def build_h2(number, title):
    p = xml_para(space_before=160, space_after=40)
    xml_run(p, f"{number}  {title}", bold=True, colour_hex="1F3864", size_pt=11)
    return [p]


def build_h3(number, title):
    p = xml_para(space_before=100, space_after=30)
    xml_run(p, f"{number}  {title}", bold=True, colour_hex="2D2D2D", size_pt=10)
    return [p]


def build_body(text, colour_hex="2D2D2D"):
    p = xml_para(space_before=40, space_after=60)
    xml_run(p, text, colour_hex=colour_hex, size_pt=10)
    return [p]


def build_bullet(text, colour_hex="2D2D2D"):
    p = xml_para("ListParagraph", space_before=20, space_after=20)
    pPr = p.find(qn("w:pPr"))
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1")
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)
    xml_run(p, text, colour_hex=colour_hex, size_pt=10)
    return [p]


def build_insert(field, hint=""):
    text = f"[INSERT: {field}]"
    if hint:
        text += f" — {hint}"
    p = xml_para(space_before=20, space_after=60)
    xml_run(p, text, bold=True, colour_hex="0059D1", size_pt=10)
    return [p]


def xml_blank():
    return xml_para(space_before=0, space_after=80)


def build_shaded_box(text, bg_hex, text_colour_hex="2D2D2D", bold=False):
    """Single-cell table used as a shaded notice box."""
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:type"), "pct"); tblW.set(qn("w:w"), "5000")
    tblPr.append(tblW)
    tbl.append(tblPr)

    tr = OxmlElement("w:tr")
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), bg_hex)
    tcPr.append(shd)
    tc.append(tcPr)

    p = xml_para(space_before=60, space_after=60)
    xml_run(p, text, bold=bold, colour_hex=text_colour_hex, size_pt=10)
    tc.append(p)
    tr.append(tc)
    tbl.append(tr)
    return [tbl, xml_blank()]


def build_table(headers, rows, hdr_bg="1F3864", suggested=False):
    """Build a full table with header row."""
    col_count = len(headers)
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle"); tblStyle.set(qn("w:val"), "TableGrid")
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:type"), "pct"); tblW.set(qn("w:w"), "5000")
    tblPr.append(tblStyle); tblPr.append(tblW)
    tbl.append(tblPr)

    def make_cell(text, bg=None, text_colour="2D2D2D", bold=False):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        if bg:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), bg)
            tcPr.append(shd)
        tc.append(tcPr)
        p = xml_para(space_before=40, space_after=40)
        # Check for blue insert placeholder
        is_insert = str(text).startswith("[INSERT:")
        is_orange = suggested and not is_insert
        cell_colour = "0059D1" if is_insert else ("C55A11" if is_orange else text_colour)
        cell_bold = bold or is_insert
        xml_run(p, str(text), bold=cell_bold, colour_hex=cell_colour, size_pt=9)
        tc.append(p)
        return tc

    # Header row
    hdr_tr = OxmlElement("w:tr")
    for h in headers:
        hdr_tr.append(make_cell(h, bg=hdr_bg, text_colour="FFFFFF", bold=True))
    tbl.append(hdr_tr)

    # Data rows
    for row in rows:
        tr = OxmlElement("w:tr")
        for i, cell_val in enumerate(row):
            tr.append(make_cell(cell_val))
        tbl.append(tr)

    return [tbl, xml_blank()]


# ── Section 0 builder ─────────────────────────────────────────────────────────

def build_section_0(brief):
    elems = []
    elems += build_shaded_box(
        "INSTRUCTION: Read, action and delete this box before this procedure is issued. "
        "Delete the entire Section 0 before this document is submitted for review or approval. "
        "It must not appear in any approved version of this procedure.",
        bg_hex="FFF2CC", text_colour_hex="C55A11", bold=True
    )
    elems += build_h1("0", "How to Use This Document")
    elems += build_body(
        "This section explains the purpose of the template, how it is structured, the sequence "
        "for completing it and the checks required before it is issued. It must be read in full "
        "before any section of this document is completed. Delete Section 0 in its entirety before issue."
    )

    elems += build_h2("0.1", "Purpose of This Template")
    elems += build_body(
        f"This template supports the preparation of a project-specific {brief['title']} for projects "
        "delivered under an Owner-Led model. Under this model, Rio Tinto directly manages all project "
        "activities, and this procedure applies to all personnel on the project, including Rio Tinto "
        "staff and all contractor organisations."
    )
    elems += build_body(
        "This template is provided as a 90% draft. Fixed content reflects the minimum Rio Tinto "
        "requirement and must not be reduced. The remaining 10% covers role names, register locations, "
        "schedules and project-specific triggers, and must be completed by the project team before "
        "this document is issued."
    )

    elems += build_h2("0.2", "How This Document is Structured")
    elems += build_body("This document contains four content types:")

    p1 = xml_para(space_before=20, space_after=20)
    xml_run(p1, "Fixed content ", bold=True, colour_hex="2D2D2D", size_pt=10)
    xml_run(p1, "is plain black text reflecting minimum Rio Tinto requirements. "
               "Fixed content must not be reduced, deleted or overridden.", colour_hex="2D2D2D", size_pt=10)
    elems.append(p1)

    p2 = xml_para(space_before=20, space_after=20)
    xml_run(p2, "[INSERT: placeholder text] ", bold=True, colour_hex="0059D1", size_pt=10)
    xml_run(p2, "appears in blue bold. Each placeholder must be replaced with accurate, "
               "project-specific information before issue.", colour_hex="2D2D2D", size_pt=10)
    elems.append(p2)

    p3 = xml_para(space_before=20, space_after=40)
    xml_run(p3, "Orange text ", colour_hex="C55A11", size_pt=10)
    xml_run(p3, "marks specific who, what, when, or how detail included as a best practice suggestion. "
               "The project team may retain, modify, or replace orange detail to suit their context.",
               colour_hex="2D2D2D", size_pt=10)
    elems.append(p3)

    elems += build_h2("0.3", "Pre-Issue Checklist")
    elems += build_body(
        "Before submitting this document for review or approval, confirm all items below are complete, "
        "then delete this checklist together with the rest of Section 0."
    )

    checklist = [
        ("Cover page is complete including document number, revision, date and signatories", "☐"),
        ("All [INSERT:] placeholder text has been replaced with project-specific information", "☐"),
        ("All instruction boxes have been read, actioned and deleted", "☐"),
        ("All role titles match the actual project organisational structure", "☐"),
        ("All system and register locations are confirmed and inserted", "☐"),
        ("All workshop and meeting dates are confirmed against the project programme", "☐"),
        ("Related procedures in the References section are issued or being issued concurrently", "☐"),
        ("Section 17 (Source Requirements Traceability) will be deleted before final issue", "☐"),
        ("Section 16 (Project Localisation Edits) has been populated with all changes", "☐"),
        ("Section 0 has been deleted in its entirety", "☐"),
    ]
    elems += build_table(
        headers=["Check", "Complete?"],
        rows=checklist,
        hdr_bg="1F3864"
    )
    return elems


# ── Roles table ───────────────────────────────────────────────────────────────

def build_roles_table():
    tbl = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle"); tblStyle.set(qn("w:val"), "TableGrid")
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:type"), "pct"); tblW.set(qn("w:w"), "5000")
    tblPr.append(tblStyle); tblPr.append(tblW)
    tbl.append(tblPr)

    def hdr_cell(text):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1F3864")
        tcPr.append(shd); tc.append(tcPr)
        p = xml_para(space_before=40, space_after=40)
        xml_run(p, text, bold=True, colour_hex="FFFFFF", size_pt=9)
        tc.append(p); return tc

    def role_group_cell(text):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D6E4F0")
        tcPr.append(shd)
        gridSpan = OxmlElement("w:gridSpan"); gridSpan.set(qn("w:val"), "4")
        tcPr.append(gridSpan)
        tc.append(tcPr)
        p = xml_para(space_before=40, space_after=40)
        xml_run(p, text, bold=True, colour_hex="1F3864", size_pt=9)
        tc.append(p); return tc

    def resp_cell(text, tick=""):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr"); tc.append(tcPr)
        p = xml_para(space_before=30, space_after=30)
        xml_run(p, text, colour_hex="2D2D2D", size_pt=9)
        tc.append(p); return tc

    def tick_cell(tick):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr"); tc.append(tcPr)
        p = xml_para(space_before=30, space_after=30)
        xml_run(p, tick, bold=True, colour_hex="2D2D2D", size_pt=9)
        tc.append(p); return tc

    def insert_cell():
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        gridSpan = OxmlElement("w:gridSpan"); gridSpan.set(qn("w:val"), "4")
        tcPr.append(gridSpan); tc.append(tcPr)
        p = xml_para(space_before=20, space_after=20)
        xml_run(p, "[INSERT: add further responsibilities as required]",
                bold=True, colour_hex="0059D1", size_pt=9)
        tc.append(p); return tc

    # Header
    hdr_tr = OxmlElement("w:tr")
    for h in ["Responsibilities", "Company", "Contractor", "Section"]:
        hdr_tr.append(hdr_cell(h))
    tbl.append(hdr_tr)

    role_groups = [
        ("Project Management", [
            ("Holds overall accountability for project compliance with this procedure and approves this procedure and any material amendment.", "✓", ""),
        ]),
        ("HSE/CSP Management", [
            ("Maintains and implements this procedure and confirms all activities meet its requirements before work proceeds.", "✓", "✓"),
            ("Reviews and approves all required assessments and plans produced under this procedure.", "✓", "✓"),
            ("Reports performance against this procedure at the agreed project HSESC performance meeting.", "✓", "✓"),
        ]),
        ("Frontline Supervision", [
            ("Ensures work crew compliance with this procedure before and during all relevant task executions.", "✓", "✓"),
            ("Stops work when the requirements of this procedure cannot be met and notifies the HSESC Manager.", "✓", "✓"),
        ]),
        ("All Project Personnel / Workers", [
            ("Complies with the requirements of this procedure in all relevant work activities.", "✓", "✓"),
            ("Reports any non-compliance with this procedure immediately to their supervisor.", "✓", "✓"),
        ]),
    ]

    for role_name, responsibilities in role_groups:
        grp_tr = OxmlElement("w:tr")
        grp_tr.append(role_group_cell(role_name))
        tbl.append(grp_tr)
        for resp_text, company, contractor in responsibilities:
            r = OxmlElement("w:tr")
            r.append(resp_cell(resp_text))
            r.append(tick_cell(company))
            r.append(tick_cell(contractor))
            r.append(tick_cell(""))
            tbl.append(r)
        ins_tr = OxmlElement("w:tr")
        ins_tr.append(insert_cell())
        tbl.append(ins_tr)

    return [tbl, xml_blank()]


# ── Content section renderer ──────────────────────────────────────────────────

def render_content_item(item):
    elems = []
    if isinstance(item, str):
        elems += build_body(item)
        return elems

    t = item.get("type", "paragraph")

    if t == "paragraph":
        colour = "C55A11" if item.get("suggested") else "2D2D2D"
        elems += build_body(item.get("text", ""), colour_hex=colour)

    elif t == "insert":
        elems += build_insert(item.get("field", "value"), item.get("hint", ""))

    elif t == "bullet_list":
        lead = item.get("lead", "")
        if lead:
            elems += build_body(lead)
        for b in item.get("items", []):
            if isinstance(b, dict):
                if b.get("insert"):
                    p = xml_para("ListParagraph", space_before=20, space_after=20)
                    pPr = p.find(qn("w:pPr"))
                    numPr = OxmlElement("w:numPr")
                    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
                    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1")
                    numPr.append(ilvl); numPr.append(numId)
                    pPr.append(numPr)
                    xml_run(p, f"[INSERT: {b.get('field','value')}]",
                            bold=True, colour_hex="0059D1", size_pt=10)
                    elems.append(p)
                else:
                    colour = "C55A11" if b.get("suggested") else "2D2D2D"
                    elems += build_bullet(b.get("text", ""), colour_hex=colour)
            else:
                elems += build_bullet(str(b))

    elif t == "table":
        headers = item.get("headers", [])
        rows = item.get("rows", [])
        caption = item.get("caption", "")
        if caption:
            elems += build_body(caption)
        if headers:
            elems += build_table(headers, rows, suggested=item.get("suggested", False))

    elif t == "instruction_box":
        elems += build_shaded_box(item.get("text", ""), bg_hex="FFF2CC",
                                   text_colour_hex="C55A11", bold=True)
    return elems


def build_content_sections(sections_data):
    elems = []
    for i, sec in enumerate(sections_data):
        sec_num = i + 3  # Sections start at 3 (1=Purpose, 2=Roles)
        title = sec.get("title", "")
        intro = sec.get("intro", "")
        subsections = sec.get("subsections", [])

        elems += build_h1(str(sec_num), title)
        if intro:
            elems += build_body(intro)

        sub_num = 1
        for sub in subsections:
            sub_title = sub.get("title", "")
            sub_intro = sub.get("intro", "")

            if sub_title:
                elems += build_h2(f"{sec_num}.{sub_num}", sub_title)
            if sub_intro:
                elems += build_body(sub_intro)

            # Level 3
            for l3 in sub.get("level3", []):
                l3_title = l3.get("title", "")
                l3_num = l3.get("number", sub_num)
                if l3_title:
                    elems += build_h3(f"{sec_num}.{sub_num}.{l3_num}", l3_title)
                for item in l3.get("content", []):
                    elems += render_content_item(item)

            for item in sub.get("content", []):
                elems += render_content_item(item)

            if sub_title:
                sub_num += 1

    return elems, sec_num + 1  # return next section number


def build_definitions(n, definitions):
    elems = []
    elems += build_h1(str(n), "Definitions and Acronyms")
    elems += build_body(
        "The following terms are used throughout this procedure. Where a term is also used in other "
        "project HSESC procedures, the definition below applies consistently across the system."
    )
    rows = [(d.get("term", ""), d.get("definition", "")) for d in definitions]
    rows.append(("[INSERT: add project-specific terms]", "[INSERT: definition]"))
    elems += build_table(["Term", "Definition"], rows)
    return elems, n + 1


def build_references(n, references):
    elems = []
    elems += build_h1(str(n), "References and Source Information")
    elems += build_body(
        "This section lists the internal and external documents this procedure relies on. "
        "This section remains in the procedure once published. Where a referenced document is "
        "updated, this procedure is reviewed to confirm the reference remains current."
    )
    rows = [(r.get("document", ""), r.get("reference", ""), r.get("type", "")) for r in references]
    rows.append(("[INSERT: further referenced document, where applicable]", "[INSERT: reference]", "[INSERT: type]"))
    elems += build_table(["Document", "Reference", "Type"], rows)
    return elems, n + 1


def build_localisation_edits(n):
    elems = []
    elems += build_h1(str(n), "Project Localisation Edits")
    elems += build_body(
        "This section captures the project-specific edits made to this document during preparation. "
        "This section stays in the document once published and is updated at each document revision. "
        "Record all additions, deletions, and wording changes from the original template."
    )
    for heading in ["Records of wording changes", "Records of new content additions", "Records of content removal"]:
        elems += build_body(heading)
        rows = [("[INSERT: description]", "[INSERT: section]")] * 3
        elems += build_table(["Description of the Change", "Section"], rows)
    return elems, n + 1


def build_traceability(n, brief):
    elems = []
    elems += build_shaded_box(
        "REMOVABLE SECTION — Source Requirements Traceability. Delete this entire section before "
        "the document is issued for approval. Table A records requirements addressed in this procedure. "
        "Table B records requirements not fully addressed, with written justification for each.",
        bg_hex="FCE4D6", text_colour_hex="C55A11", bold=True
    )
    elems += build_h1(str(n), "Source Requirements Traceability (DELETE BEFORE ISSUE)")

    elems += build_h2(f"{n}.1", "Requirements Addressed in This Procedure")
    rows_a = []
    for src in brief.get("external_sources", []):
        rows_a.append((
            "[INSERT: section]",
            f"{src.get('framework','')} {src.get('clause','')}".strip(),
            src.get("summary", ""),
            "Full"
        ))
    if not rows_a:
        rows_a = [("[INSERT: section]", "[INSERT: source]", "[INSERT: requirement]", "[INSERT: status]")]
    elems += build_table(["Section", "Source", "Requirement", "Status"], rows_a)

    elems += build_h2(f"{n}.2", "Requirements Not Fully Addressed in This Procedure")
    elems += build_body(
        "List any requirements from the external sources above that are not fully addressed in this "
        "document, with a specific written justification for each."
    )
    rows_b = [("[INSERT: description of requirement not covered]",
               "[INSERT: source and clause]",
               "[INSERT: specific written justification]")]
    elems += build_table(["Requirement Not Covered", "Source", "Justification"], rows_b)
    return elems


# ── Template injection ────────────────────────────────────────────────────────

def inject_into_template(template_path, output_path, brief, all_elements):
    """
    Copy the template docx, then inject all_elements into document.xml
    by replacing the 'Enter Body Content Text' placeholder paragraph.
    Also replaces PROCEDURE TITLE with the actual procedure title.
    """
    shutil.copy2(template_path, output_path)

    # Work directly on the zip
    import tempfile
    tmp_dir = tempfile.mkdtemp()

    with zipfile.ZipFile(output_path, 'r') as z:
        z.extractall(tmp_dir)

    doc_path = os.path.join(tmp_dir, "word", "document.xml")

    tree = etree.parse(doc_path)
    root = tree.getroot()
    ns = {"w": W}

    body = root.find(".//w:body", ns)

    # 1. Replace PROCEDURE TITLE text nodes
    for t_elem in root.iter(qn("w:t")):
        if t_elem.text and "PROCEDURE TITLE" in t_elem.text:
            proc_title = brief["title"].replace(" Procedure", "").replace(" Programme", "")
            t_elem.text = t_elem.text.replace("PROCEDURE TITLE", proc_title)

    # 2. Find the placeholder paragraph and replace it with all generated elements
    placeholder_para = None
    for p in body.findall("w:p", ns):
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if "Enter Body Content Text" in text:
            placeholder_para = p
            break

    if placeholder_para is None:
        print("  WARNING: Could not find body content placeholder — appending before sectPr")
        # Fall back: insert before sectPr
        sectPr = body.find("w:sectPr", ns)
        insert_idx = list(body).index(sectPr) if sectPr is not None else len(list(body))
        for i, elem in enumerate(all_elements):
            body.insert(insert_idx + i, elem)
    else:
        # Insert all elements in place of the placeholder
        idx = list(body).index(placeholder_para)
        body.remove(placeholder_para)
        for i, elem in enumerate(all_elements):
            body.insert(idx + i, elem)

    # Write back
    tree.write(doc_path, xml_declaration=True, encoding="UTF-8", standalone=True)

    # Rezip
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for root_dir, dirs, files in os.walk(tmp_dir):
            for file in files:
                file_path = os.path.join(root_dir, file)
                arcname = os.path.relpath(file_path, tmp_dir)
                zout.write(file_path, arcname)

    shutil.rmtree(tmp_dir)
    print(f"  Document written on template: {output_path}")


# ── Claude generation ─────────────────────────────────────────────────────────

def _call_claude(client, system, prompt, label=""):
    for attempt in range(3):
        try:
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=16000,
                system=system,
                messages=[{"role": "user", "content": prompt}],
            )
            raw = response.content[0].text.strip()
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
            return json.loads(raw)
        except json.JSONDecodeError as e:
            print(f"  JSON error {label} attempt {attempt+1}: {e}")
            if attempt < 2:
                time.sleep(5)
            else:
                raise
        except Exception as e:
            print(f"  API error {label} attempt {attempt+1}: {e}")
            if attempt < 2:
                time.sleep(10)
            else:
                raise


def generate_content(client, brief, example_text):
    requirements_text = "\n".join(
        f"- [{r.get('tag','RT Standard')}] under '{r.get('section_heading','')}': {r.get('text','')}"
        for r in brief.get("requirements", [])
    )
    pcg_text = "\n".join(
        f"{i+1}. {item}"
        for i, item in enumerate(brief.get("project_completion_guidance", []))
    )
    owns_text = "\n".join(f"- {o}" for o in brief.get("owns", []))
    refs_text = "\n".join(f"- {r}" for r in brief.get("references", []))

    system = (
        "You are a document engineering assistant generating HSESC management system procedures "
        "for Rio Tinto owner-led projects. Output only valid JSON — no markdown, no code fences."
    )

    style_block = f"""STYLE REFERENCE — match this example procedure exactly in structure, heading depth,
sentence patterns, paragraph length, table formats, INSERT placeholder usage, and content density:
Key patterns:
1. Level-1 sections: 25-35 word intro. Level-2: 15-25 word intro.
2. Paragraphs 2-4 sentences. Lead-in sentences end with a colon before bullet lists.
3. Tables pre-filled: fixed rows for standard content, [INSERT:] rows for project data.
4. INSERT placeholders embedded inside sentences, never standalone.
5. Orange suggested content gives a concrete default the project may keep or change.
6. Impersonal voice — subject is always a role, system, or document.

EXAMPLE:
{example_text}"""

    content_spec = """Content item types for 'content' arrays:
{"type":"paragraph","text":"<text>","suggested":false}  (suggested:true for orange text)
{"type":"insert","field":"<name>","hint":"<short hint>"}
{"type":"bullet_list","lead":"<lead-in ending with colon>","items":["<text>",{"text":"<suggested>","suggested":true},{"insert":true,"field":"<name>"}]}
{"type":"table","caption":"<optional>","suggested":false,"headers":["<col>"],"rows":[["<cell>"],["[INSERT: value]"]]}
{"type":"instruction_box","text":"<text>"}"""

    # Call 1: sections
    print("  Call 1/2: generating body sections...")
    sections_prompt = f"""Generate body sections for:
PROCEDURE: {brief['procedure_id']} — {brief['title']}
Document: {brief['document_number']} | Element: {brief['management_system_element']}

OWNS: {owns_text}
REFERENCES (cross-reference only, do not repeat): {refs_text}
REQUIREMENTS (all must appear): {requirements_text}
PROJECT COMPLETION GUIDANCE (embed as [INSERT:] in relevant sections): {pcg_text}

{style_block}
{content_spec}

Return ONLY this JSON:
{{"sections":[{{"title":"<Level-1 title>","intro":"<25-35 words>","subsections":[{{"title":"<Level-2 or empty>","intro":"<15-25 words or empty>","level3":[{{"title":"<Level-3>","number":<int>,"content":[<items>]}}],"content":[<items>]}}]}}]}}

RULES: Every requirement must appear. Do not reproduce content owned by referenced documents.
Standard PTW phrase: "raise a permit via the site PTW system before commencing work, refer RTPR-HSE-PRO-0021."
Standard incident phrase: "report all incidents and near misses per project requirements, refer RTPR-HSE-PRO-0009."
Return only valid JSON."""

    sections_result = _call_claude(client, system, sections_prompt, "sections")

    # Call 2: definitions + references
    print("  Call 2/2: generating definitions and references...")
    section_titles = [s.get("title","") for s in sections_result.get("sections",[])]

    defsrefs_prompt = f"""Generate definitions and references for:
PROCEDURE: {brief['procedure_id']} — {brief['title']}
Body sections cover: {', '.join(section_titles)}
Referenced documents: {refs_text}

Return ONLY this JSON:
{{"definitions":[{{"term":"<acronym>","definition":"<definition>"}}],"references":[{{"document":"<title>","reference":"<number>","type":"<Internal procedure|External RT|External framework>"}}]}}

Include definitions for all acronyms used. Include all cross-referenced procedures in references.
Always include RT Risk Management Standard and Appendix C C-1 in references.
Return only valid JSON."""

    defsrefs_result = _call_claude(client, system, defsrefs_prompt, "defs+refs")

    return {
        "sections": sections_result.get("sections", []),
        "definitions": defsrefs_result.get("definitions", []),
        "references": defsrefs_result.get("references", []),
    }


# ── Main ──────────────────────────────────────────────────────────────────────

def read_docx_text(path):
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in parts:
                    parts.append(t)
    return "\n".join(parts)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--procedure-id",  required=True)
    parser.add_argument("--brief-dir",     required=True)
    parser.add_argument("--template-dir",  required=True)
    parser.add_argument("--output-dir",    required=True)
    parser.add_argument("--inputs-dir",    default="inputs")
    args = parser.parse_args()

    brief_path = os.path.join(args.brief_dir, f"{args.procedure_id}.json")
    if not os.path.exists(brief_path):
        print(f"ERROR: Brief not found: {brief_path}")
        sys.exit(1)

    with open(brief_path, "r", encoding="utf-8") as f:
        brief = json.load(f)

    # Find template
    template_path = os.path.join(args.inputs_dir, "procedure_template.docx")
    if not os.path.exists(template_path):
        template_path = os.path.join(args.template_dir, "procedure_template.docx")
    if not os.path.exists(template_path):
        print("ERROR: procedure_template.docx not found in inputs/ or template/")
        sys.exit(1)
    print(f"  Using template: {template_path}")

    # Find example
    example_path = os.path.join(args.inputs_dir, "example_procedure.docx")
    if os.path.exists(example_path):
        print(f"  Using style reference: example_procedure.docx")
        example_text = read_docx_text(example_path)
        if len(example_text) > 10000:
            example_text = example_text[:10000] + "\n[... truncated ...]"
    else:
        print("  WARNING: No example_procedure.docx found.")
        example_text = "(No example provided.)"

    os.makedirs(args.output_dir, exist_ok=True)

    safe = re.sub(r"[^\w\s-]", "", brief["title"].replace(" Procedure","").replace(" Programme",""))
    safe = re.sub(r"\s+", "_", safe.strip())
    output_path = os.path.join(args.output_dir, f"{args.procedure_id}_{safe}.docx")

    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    print(f"Generating content for {args.procedure_id}: {brief['title']}...")
    generated = generate_content(client, brief, example_text)

    print("Assembling document on template...")

    # Build all content elements
    all_elements = []

    # Section 0
    all_elements += build_section_0(brief)

    # TOC placeholder note
    all_elements += build_shaded_box(
        "Update the table of contents: right-click the Contents page and select 'Update Field' "
        "after all placeholder content has been completed.",
        bg_hex="DEEAF1", text_colour_hex="1F3864"
    )

    # Section 1 — Purpose and Scope
    all_elements += build_h1("1", "Purpose and Scope")
    all_elements += build_body(
        "This section defines the purpose of this procedure, the project lifecycle phases within "
        "scope, and the obligations it establishes from project Execution through to Handover."
    )
    all_elements += build_h2("1.1", "Purpose of This Document")
    all_elements += build_body(
        "This section states the purpose of this procedure and the requirements it establishes for the project."
    )
    preamble = brief.get("preamble", "")
    if preamble:
        all_elements += build_body(preamble)
    else:
        all_elements += build_insert("purpose statement",
            "describe what this procedure governs and the project context it applies to")

    all_elements += build_h2("1.2", "Scope of This Document")
    all_elements += build_body(
        "This section defines which activities and lifecycle phases are governed by this document."
    )
    all_elements += build_body(
        f"This procedure applies to all Rio Tinto project personnel and all contractor organisations "
        f"performing work under an Owner-Led project. It governs "
        f"{brief['title'].replace(' Procedure','').replace(' Programme','').lower()} "
        f"from project Execution through to Handover."
    )
    all_elements += build_insert("scope qualifications",
        "note any activities, phases, or locations excluded from this procedure's scope")

    # Section 2 — Roles and Responsibilities
    all_elements += build_h1("2", "Roles and Responsibilities")
    all_elements += build_body(
        "This section defines the responsibilities for all project roles relevant to this procedure, "
        "covering the Rio Tinto project team and all contractor organisations. Responsibilities are "
        "aligned to governance, oversight, assessment execution, reporting, and control verification."
    )
    all_elements += build_roles_table()

    # Content sections (3 onwards)
    content_elems, next_n = build_content_sections(generated.get("sections", []))
    all_elements += content_elems

    # Back matter
    defs_elems, next_n = build_definitions(next_n, generated.get("definitions", []))
    all_elements += defs_elems

    refs_elems, next_n = build_references(next_n, generated.get("references", []))
    all_elements += refs_elems

    loc_elems, next_n = build_localisation_edits(next_n)
    all_elements += loc_elems

    all_elements += build_traceability(next_n, brief)

    inject_into_template(template_path, output_path, brief, all_elements)
    print(f"Done: {os.path.basename(output_path)}")


if __name__ == "__main__":
    main()
